"""Tests for scripts/generate_report.py — Word document report generation."""

import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SCRIPT = str(PROJECT_ROOT / "scripts" / "generate_report.py")
FIXTURE_FILE = PROJECT_ROOT / "tests" / "fixtures" / "sample-feature.json"


def _run_script(*args: str, cwd: Path = PROJECT_ROOT) -> subprocess.CompletedProcess:
    """Run generate_report.py with the given CLI args."""
    return subprocess.run(
        [sys.executable, SCRIPT, *args],
        capture_output=True,
        text=True,
        cwd=str(cwd),
    )


def _write_input_json(data: dict, directory: Path) -> Path:
    """Serialize *data* to a JSON file inside *directory* and return the path."""
    path = directory / "input.json"
    path.write_text(json.dumps(data, default=str))
    return path


def _ensure_sources_unavailable(data: dict) -> dict:
    """Add ``sources_unavailable`` if missing so the JSON validates against the model."""
    if "sources_unavailable" not in data:
        data = {**data, "sources_unavailable": []}
    return data


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestReportFromSampleFixture:
    """Generate a report from the sample-feature.json fixture file."""

    def test_report_from_sample_fixture(self, tmp_output_dir: Path):
        output_path = tmp_output_dir / "report.docx"
        result = _run_script(
            "--input", str(FIXTURE_FILE),
            "--output", str(output_path),
        )

        assert result.returncode == 0, (
            f"Script failed (rc={result.returncode}):\n{result.stderr}"
        )
        assert output_path.exists()
        assert output_path.stat().st_size > 0


class TestReportContainsRequiredSections:
    """Verify the generated .docx includes all required section headings."""

    REQUIRED_SECTIONS = [
        "Executive Summary",
        "Feature Overview",
        "Business Value Analysis",
        "Leading Indicators",
        "Strategic Value",
        "Assumptions",
        "Source Citations",
    ]

    def test_report_contains_required_sections(
        self, sample_report_input: dict, tmp_output_dir: Path, tmp_path: Path,
    ):
        data = _ensure_sources_unavailable(sample_report_input)
        input_path = _write_input_json(data, tmp_path)
        output_path = tmp_output_dir / "report.docx"

        result = _run_script(
            "--input", str(input_path),
            "--output", str(output_path),
        )
        assert result.returncode == 0, (
            f"Script failed (rc={result.returncode}):\n{result.stderr}"
        )

        doc = Document(str(output_path))
        headings = [
            p.text
            for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]

        for section in self.REQUIRED_SECTIONS:
            assert any(section in h for h in headings), (
                f"Missing heading containing '{section}'. Found headings: {headings}"
            )


class TestChartEmbedding:
    """Verify that chart images are embedded when --charts-dir is provided."""

    def test_chart_embedding(
        self, sample_report_input: dict, tmp_output_dir: Path, tmp_path: Path,
    ):
        import matplotlib.pyplot as plt

        charts_dir = tmp_path / "charts"
        charts_dir.mkdir()
        chart_path = charts_dir / "test_chart.png"

        fig, ax = plt.subplots(figsize=(2, 2))
        ax.text(0.5, 0.5, "Test", ha="center")
        fig.savefig(str(chart_path), dpi=100)
        plt.close(fig)

        data = _ensure_sources_unavailable(sample_report_input)
        input_path = _write_input_json(data, tmp_path)
        output_path = tmp_output_dir / "report.docx"

        result = _run_script(
            "--input", str(input_path),
            "--output", str(output_path),
            "--charts-dir", str(charts_dir),
        )
        assert result.returncode == 0, (
            f"Script failed (rc={result.returncode}):\n{result.stderr}"
        )

        doc = Document(str(output_path))
        inline_shapes = doc.inline_shapes
        assert len(inline_shapes) >= 1, "Expected at least one embedded image (chart)"


class TestMissingTemplateFallback:
    """The script should fall back to a plain document when the template is missing."""

    def test_missing_template_fallback(
        self, sample_report_input: dict, tmp_output_dir: Path, tmp_path: Path,
    ):
        data = _ensure_sources_unavailable(sample_report_input)
        input_path = _write_input_json(data, tmp_path)
        output_path = tmp_output_dir / "report.docx"

        result = _run_script(
            "--input", str(input_path),
            "--output", str(output_path),
            "--template", str(tmp_path / "nonexistent.dotx"),
        )
        assert result.returncode == 0, (
            f"Expected fallback to plain doc, but script failed "
            f"(rc={result.returncode}):\n{result.stderr}"
        )
        assert output_path.exists()


class TestOutputFilenameAutogeneration:
    """When --output is omitted the script should auto-generate business-case-*.docx."""

    def test_output_filename_autogeneration(
        self, sample_report_input: dict, tmp_path: Path,
    ):
        data = _ensure_sources_unavailable(sample_report_input)
        input_path = _write_input_json(data, tmp_path)

        result = _run_script("--input", str(input_path), cwd=tmp_path)
        assert result.returncode == 0, (
            f"Script failed (rc={result.returncode}):\n{result.stderr}"
        )

        generated = list(tmp_path.glob("business-case-*.docx"))
        assert generated, (
            f"No business-case-*.docx found in {tmp_path}. "
            f"Files present: {list(tmp_path.iterdir())}"
        )


class TestInvalidInput:
    """The script should exit with code 1 on bad input."""

    def test_invalid_json_exits_with_error(self, tmp_path: Path):
        bad_file = tmp_path / "bad.json"
        bad_file.write_text("{this is not valid json!!!")

        result = _run_script("--input", str(bad_file))

        assert result.returncode == 1
        assert result.stderr.strip(), "Expected an error message on stderr"

    def test_validation_error_exits(self, tmp_path: Path):
        incomplete = tmp_path / "incomplete.json"
        incomplete.write_text(json.dumps({"not_a_valid_field": True}))

        result = _run_script("--input", str(incomplete))

        assert result.returncode == 1
