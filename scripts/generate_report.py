"""Assemble a Word document (.docx) business-case report from a ReportInput JSON file."""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pydantic import ValidationError

# Allow imports from the project root
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from scripts.models import Confidence, Reliability, ReportInput


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def format_dollar(value: float) -> str:
    if abs(value) >= 1_000_000:
        return f"${value / 1_000_000:.1f}M"
    elif abs(value) >= 1_000:
        return f"${value / 1_000:.0f}k"
    else:
        return f"${value:.0f}"


def _confidence_badge(confidence: Confidence) -> str:
    return {
        Confidence.weak: "[WEAK]",
        Confidence.medium: "[MEDIUM]",
        Confidence.strong: "[STRONG]",
    }[confidence]


def _sort_key_value_high(est) -> tuple[int, float]:
    """Sort key: estimates with a dollar_value_high come first (descending), nulls last."""
    if est.dollar_value_high is None:
        return (1, 0.0)
    return (0, -est.dollar_value_high)


def _auto_output_path(title: str) -> str:
    sanitized = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
    date_str = datetime.now().strftime("%Y-%m-%d")
    return f"business-case-{sanitized}-{date_str}.docx"


def _dollar_unit_label(unit_value: str) -> str:
    return {"one_time": "one-time", "per_year": "per year", "per_month": "per month"}.get(
        unit_value, unit_value
    )


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, report: ReportInput) -> None:
    doc.add_heading(report.feature.title, level=0)
    doc.add_paragraph("Business Case Report", style="Subtitle")
    doc.add_paragraph(f"Date: {report.generated_at}")
    mode_label = (
        "Interactive Analysis" if report.mode.value == "interactive" else "Autonomous Analysis"
    )
    doc.add_paragraph(f"Mode: {mode_label}")
    doc.add_page_break()


def _add_executive_summary(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.strategic_value.summary)

    sorted_estimates = sorted(report.value_estimates, key=_sort_key_value_high)
    top = [e for e in sorted_estimates if e.dollar_value_high is not None][:3]
    if top:
        doc.add_paragraph("Top value drivers:")
        for est in top:
            low = format_dollar(est.dollar_value_low) if est.dollar_value_low is not None else "N/A"
            high = format_dollar(est.dollar_value_high)
            doc.add_paragraph(
                f"  • {est.category_name}: {low} – {high} ({_dollar_unit_label(est.dollar_unit.value)})",
            )

    # Total value range
    lows = [e.dollar_value_low for e in report.value_estimates if e.dollar_value_low is not None]
    highs = [e.dollar_value_high for e in report.value_estimates if e.dollar_value_high is not None]
    if lows or highs:
        total_low = format_dollar(sum(lows)) if lows else "N/A"
        total_high = format_dollar(sum(highs)) if highs else "N/A"
        p = doc.add_paragraph()
        p.add_run("Total estimated value range: ").bold = True
        p.add_run(f"{total_low} – {total_high}")


def _add_feature_overview(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Feature Overview", level=1)
    doc.add_heading("Description", level=2)
    doc.add_paragraph(report.feature.description)

    doc.add_heading("Target Users", level=2)
    for user in report.feature.target_users:
        doc.add_paragraph(user, style="List Bullet")

    doc.add_heading("Source", level=2)
    source_text = f"Format: {report.feature.source_format.value}"
    if report.feature.source_reference:
        source_text += f"\nReference: {report.feature.source_reference}"
    doc.add_paragraph(source_text)


def _build_citation_index(report: ReportInput) -> dict[str, int]:
    """Build a mapping from citation ID to sequential footnote number."""
    seen: set[str] = set()
    index: dict[str, int] = {}
    num = 1
    for est in sorted(report.value_estimates, key=_sort_key_value_high):
        for cite in est.citations:
            if cite.id not in seen:
                seen.add(cite.id)
                index[cite.id] = num
                num += 1
    return index


def _reliability_badge(reliability: Reliability) -> str:
    return {
        Reliability.primary: "\u2605\u2605\u2605 Primary",
        Reliability.secondary: "\u2605\u2605 Secondary",
        Reliability.tertiary: "\u2605 Tertiary",
    }[reliability]


def _add_business_value_analysis(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Business Value Analysis", level=1)

    citation_index = _build_citation_index(report)
    sorted_estimates = sorted(report.value_estimates, key=_sort_key_value_high)

    for est in sorted_estimates:
        doc.add_heading(est.category_name, level=2)
        if est.subcategory:
            doc.add_heading(est.subcategory, level=3)

        # Value range
        p = doc.add_paragraph()
        p.add_run("Value Range: ").bold = True
        if est.dollar_value_low is not None and est.dollar_value_high is not None:
            run_low = p.add_run(format_dollar(est.dollar_value_low))
            run_low.bold = True
            p.add_run(" – ")
            run_high = p.add_run(format_dollar(est.dollar_value_high))
            run_high.bold = True
            p.add_run(f" {_dollar_unit_label(est.dollar_unit.value)}")
        else:
            p.add_run("Not quantified")

        # Confidence
        p = doc.add_paragraph()
        p.add_run("Confidence: ").bold = True
        p.add_run(f"{est.confidence.value} {_confidence_badge(est.confidence)}")

        # Hard dollar
        p = doc.add_paragraph()
        p.add_run("Hard Dollar: ").bold = True
        p.add_run("Yes" if est.is_hard_dollar else "No")

        # Reasoning
        doc.add_paragraph(est.reasoning)

        # Methodology
        doc.add_heading("Methodology", level=3)
        doc.add_paragraph(est.methodology)

        # Assumptions
        doc.add_heading("Assumptions", level=3)
        for assumption in est.assumptions:
            doc.add_paragraph(assumption, style="List Bullet")

        # Sources / citations — numbered footnote-style references
        doc.add_heading("Sources", level=3)
        for cite in est.citations:
            fn = citation_index.get(cite.id, 0)
            label = f"[{fn}] {cite.source_type.value} — {cite.title}"
            doc.add_paragraph(label)
            excerpt_para = doc.add_paragraph(cite.excerpt)
            excerpt_para.paragraph_format.left_indent = Inches(0.5)
            if cite.url:
                url_para = doc.add_paragraph(cite.url)
                url_para.paragraph_format.left_indent = Inches(0.5)


def _add_charts(doc: Document, charts_dir: Path) -> None:
    pngs = sorted(charts_dir.glob("*.png"))
    if not pngs:
        return
    doc.add_heading("Visual Analysis", level=1)
    for png_path in pngs:
        try:
            doc.add_picture(str(png_path), width=Inches(6))
            caption = png_path.stem.replace("_", " ").replace("-", " ").title()
            cap_para = doc.add_paragraph(caption)
            cap_para.alignment = 1  # center
        except Exception as exc:
            print(f"Warning: could not embed chart {png_path}: {exc}", file=sys.stderr)


def _add_leading_indicators(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Leading Indicators", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Indicator", "Linked Outcome", "Expected Impact"]):
        hdr[idx].text = text
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for li in report.leading_indicators:
        row = table.add_row().cells
        row[0].text = li.indicator
        row[1].text = li.linked_outcome
        row[2].text = li.expected_impact


def _add_strategic_value(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Strategic Value Assessment", level=1)
    doc.add_paragraph(report.strategic_value.summary)
    doc.add_heading("Strategic Dimensions", level=2)
    for dim in report.strategic_value.dimensions:
        doc.add_paragraph(dim, style="List Bullet")


def _add_assumptions_methodology(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Assumptions and Methodology", level=1)
    for assumption in report.assumptions_made:
        doc.add_paragraph(assumption, style="List Bullet")

    if report.sources_unavailable:
        doc.add_heading("Sources Unavailable", level=2)
        for src in report.sources_unavailable:
            doc.add_paragraph(src, style="List Bullet")


def _add_source_citations(doc: Document, report: ReportInput) -> None:
    doc.add_heading("Source Citations", level=1)

    citation_index = _build_citation_index(report)

    # Deduplicate citations across all value estimates (in footnote order)
    seen: set[str] = set()
    all_citations = []
    for est in sorted(report.value_estimates, key=_sort_key_value_high):
        for cite in est.citations:
            if cite.id not in seen:
                seen.add(cite.id)
                all_citations.append(cite)

    if not all_citations:
        doc.add_paragraph("No source citations available.")
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["#", "Source Type", "Title", "Reliability", "Excerpt"]):
        hdr[idx].text = text
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for cite in all_citations:
        fn = citation_index.get(cite.id, 0)
        row = table.add_row().cells
        row[0].text = str(fn)
        row[1].text = cite.source_type.value
        title_text = cite.title
        if cite.url:
            title_text += f"\n{cite.url}"
        row[2].text = title_text
        row[3].text = _reliability_badge(cite.reliability)
        row[4].text = cite.excerpt

    # Sources Searched appendix
    if report.sources_unavailable:
        doc.add_heading("Sources Searched", level=2)
        doc.add_paragraph(
            "The following sources were searched but did not return usable results:"
        )
        for src in report.sources_unavailable:
            doc.add_paragraph(src, style="List Bullet")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_report(report: ReportInput, template_path: Path | None, charts_dir: Path | None) -> Document:
    # Try to open template
    doc = None
    if template_path and template_path.exists():
        try:
            doc = Document(str(template_path))
        except Exception as exc:
            print(f"Warning: could not open template {template_path}: {exc}", file=sys.stderr)
    if doc is None:
        if template_path and not template_path.exists():
            print(f"Warning: template not found at {template_path}, using plain document", file=sys.stderr)
        doc = Document()

    _add_cover_page(doc, report)
    _add_executive_summary(doc, report)

    # Autonomous mode: surface assumptions prominently before the main analysis
    if report.mode.value == "autonomous" and report.assumptions_made:
        doc.add_heading("Assumptions and Agent Decisions", level=1)
        doc.add_paragraph(
            "This report was generated in autonomous mode without human review. "
            "The following assumptions were made during analysis:"
        )
        for assumption in report.assumptions_made:
            doc.add_paragraph(assumption, style="List Bullet")
        doc.add_paragraph(
            "Review these assumptions carefully. Estimates may change significantly "
            "if any assumption proves incorrect."
        )

    _add_feature_overview(doc, report)
    _add_business_value_analysis(doc, report)

    if charts_dir and charts_dir.is_dir():
        _add_charts(doc, charts_dir)

    _add_leading_indicators(doc, report)
    _add_strategic_value(doc, report)
    _add_assumptions_methodology(doc, report)
    _add_source_citations(doc, report)

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Word document business-case report from a ReportInput JSON file."
    )
    parser.add_argument("--input", required=True, help="Path to ReportInput JSON file")
    parser.add_argument("--output", default=None, help="Output path for .docx file")
    parser.add_argument(
        "--template",
        default=str(Path(__file__).resolve().parent.parent / "templates" / "business-case-template.dotx"),
        help="Path to .dotx template (default: templates/business-case-template.dotx)",
    )
    parser.add_argument("--charts-dir", default=None, help="Directory containing chart PNGs to embed")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Parse JSON
    try:
        with open(input_path) as f:
            raw = json.load(f)
    except json.JSONDecodeError as exc:
        print(f"Error: invalid JSON: {exc}", file=sys.stderr)
        sys.exit(1)

    # Validate with Pydantic
    try:
        report = ReportInput(**raw)
    except ValidationError as exc:
        print(f"Error: validation failed:\n{exc}", file=sys.stderr)
        sys.exit(1)

    # Resolve output path
    output_path = args.output
    if not output_path:
        output_path = _auto_output_path(report.feature.title)

    template_path = Path(args.template) if args.template else None
    charts_dir = Path(args.charts_dir) if args.charts_dir else None

    doc = build_report(report, template_path, charts_dir)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
